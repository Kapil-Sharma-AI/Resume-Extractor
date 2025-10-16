UI Code:-

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import threading
from pathlib import Path
import pandas as pd
from enhanced_candidate_extractor import EnhancedCandidateExtractor
from datetime import datetime
import json

class EnhancedCandidateExtractorGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Enhanced Candidate Resume Extractor")
        self.root.geometry("900x700")
        
        # Initialize variables
        self.folder_path = tk.StringVar()
        self.output_format = tk.StringVar(value="excel")
        self.processing = False
        
        # Initialize extractor
        self.extractor = EnhancedCandidateExtractor()
        
        # Setup UI
        self.setup_ui()
        
    def setup_ui(self):
        """Setup the user interface"""
        # Create main frame
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        
        # Title
        title_label = ttk.Label(main_frame, text="Enhanced Candidate Resume Extractor", 
                               font=('Arial', 16, 'bold'))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 20))
        
        # Folder selection
        ttk.Label(main_frame, text="Select Folder:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.folder_entry = ttk.Entry(main_frame, textvariable=self.folder_path, width=50)
        self.folder_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=5, padx=(5, 0))
        
        ttk.Button(main_frame, text="Browse", 
                  command=self.browse_folder).grid(row=1, column=2, pady=5, padx=(5, 0))
        
        # Output format selection
        ttk.Label(main_frame, text="Output Format:").grid(row=2, column=0, sticky=tk.W, pady=5)
        format_frame = ttk.Frame(main_frame)
        format_frame.grid(row=2, column=1, sticky=tk.W, pady=5)
        
        ttk.Radiobutton(format_frame, text="Excel", variable=self.output_format, 
                       value="excel").pack(side=tk.LEFT, padx=(0, 10))
        ttk.Radiobutton(format_frame, text="CSV", variable=self.output_format, 
                       value="csv").pack(side=tk.LEFT, padx=(0, 10))
        ttk.Radiobutton(format_frame, text="JSON", variable=self.output_format, 
                       value="json").pack(side=tk.LEFT)
        
        # Process button
        self.process_btn = ttk.Button(main_frame, text="Process Files", 
                                     command=self.start_processing)
        self.process_btn.grid(row=3, column=0, columnspan=3, pady=20)
        
        # Progress bar
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(main_frame, variable=self.progress_var, 
                                          maximum=100, length=400)
        self.progress_bar.grid(row=4, column=0, columnspan=3, pady=5, sticky=(tk.W, tk.E))
        
        # Status label
        self.status_label = ttk.Label(main_frame, text="Ready to process files")
        self.status_label.grid(row=5, column=0, columnspan=3, pady=5)
        
        # Processing log frame
        log_frame = ttk.LabelFrame(main_frame, text="Processing Log", padding="5")
        log_frame.grid(row=6, column=0, columnspan=3, pady=(20, 0), sticky=(tk.W, tk.E, tk.N, tk.S))
        log_frame.columnconfigure(0, weight=1)
        log_frame.rowconfigure(0, weight=1)
        
        # Log text area
        self.log_text = scrolledtext.ScrolledText(log_frame, height=15, width=80)
        self.log_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Results preview frame
        results_frame = ttk.LabelFrame(main_frame, text="Results Preview", padding="5")
        results_frame.grid(row=7, column=0, columnspan=3, pady=(10, 0), sticky=(tk.W, tk.E, tk.N, tk.S))
        results_frame.columnconfigure(0, weight=1)
        results_frame.rowconfigure(0, weight=1)
        
        # Results tree
        self.results_tree = ttk.Treeview(results_frame, height=8)
        self.results_tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Scrollbar for results tree
        tree_scrollbar = ttk.Scrollbar(results_frame, orient=tk.VERTICAL, 
                                      command=self.results_tree.yview)
        tree_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        self.results_tree.configure(yscrollcommand=tree_scrollbar.set)
        
        # Configure grid weights for main frame
        main_frame.rowconfigure(6, weight=1)
        main_frame.rowconfigure(7, weight=1)
        
        # Clear log button
        ttk.Button(main_frame, text="Clear Log", 
                  command=self.clear_log).grid(row=8, column=0, pady=5, sticky=tk.W)
        
        # Export summary button
        ttk.Button(main_frame, text="Export Processing Summary", 
                  command=self.export_summary).grid(row=8, column=2, pady=5, sticky=tk.E)
    
    def browse_folder(self):
        """Browse for folder containing resume files"""
        folder = filedialog.askdirectory()
        if folder:
            self.folder_path.set(folder)
    
    def log_message(self, message: str, level: str = "info"):
        """Add message to log"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        
        # Color coding for different levels
        if level == "error":
            color = "red"
        elif level == "warning":
            color = "orange"
        elif level == "success":
            color = "green"
        else:
            color = "black"
        
        # Insert message with color
        self.log_text.insert(tk.END, f"[{timestamp}] {message}\n")
        
        # Scroll to bottom
        self.log_text.see(tk.END)
        
        # Update UI
        self.root.update_idletasks()
    
    def clear_log(self):
        """Clear the log text area"""
        self.log_text.delete(1.0, tk.END)
    
    def start_processing(self):
        """Start processing files in a separate thread"""
        if self.processing:
            return
        
        folder = self.folder_path.get()
        if not folder or not os.path.exists(folder):
            messagebox.showerror("Error", "Please select a valid folder")
            return
        
        # Disable button
        self.process_btn.config(state="disabled")
        self.processing = True
        
        # Clear previous results
        self.clear_log()
        self.clear_results_tree()
        
        # Start processing in separate thread
        thread = threading.Thread(target=self.process_files_thread, args=(folder,))
        thread.daemon = True
        thread.start()
    
    def process_files_thread(self, folder_path: str):
        """Process files in separate thread"""
        try:
            self.log_message("=== Starting Resume Processing ===", "info")
            self.update_status("Scanning for files...")
            
            # Find all supported files
            supported_extensions = ['.pdf', '.docx', '.doc']
            file_paths = []
            
            for ext in supported_extensions:
                pattern = f"*{ext}"
                files = list(Path(folder_path).glob(pattern))
                file_paths.extend(files)
            
            if not file_paths:
                self.log_message("No supported files found in the selected folder", "warning")
                self.update_status("No files to process")
                return
            
            self.log_message(f"Found {len(file_paths)} files to process", "info")
            
            # Update progress
            self.progress_var.set(0)
            
            # Process files
            results = []
            total_files = len(file_paths)
            
            for i, file_path in enumerate(file_paths):
                try:
                    self.update_status(f"Processing {file_path.name}...")
                    self.log_message(f"Processing file {i+1}/{total_files}: {file_path.name}", "info")
                    
                    # Process single file
                    result = self.extractor.process_single_file(str(file_path))
                    results.append(result)
                    
                    # Update progress
                    progress = ((i + 1) / total_files) * 100
                    self.progress_var.set(progress)
                    
                    # Log result
                    if "Error" not in result['Candidate Name'] and "Failed" not in result['Candidate Name']:
                        self.log_message(f"✓ Successfully processed: {result['Candidate Name']}", "success")
                    else:
                        self.log_message(f"✗ Failed to process: {file_path.name}", "error")
                    
                except Exception as e:
                    self.log_message(f"✗ Error processing {file_path.name}: {str(e)}", "error")
                    results.append({
                        'Filename': file_path.name,
                        'Candidate Name': 'Processing Error',
                        'Email ID': 'Processing Error',
                        'SAP Skills': 'Processing Error',
                        'Total Experience (Years)': 0,
                        'Summary': f'Error: {str(e)}',
                        'Certifications': 'Processing Error'
                    })
            
            # Create DataFrame
            df = pd.DataFrame(results)
            
            # Save results
            self.save_results(df, folder_path)
            
            # Update UI with results
            self.root.after(0, self.update_results_display, df)
            
            # Show processing summary
            self.show_processing_summary(df)
            
        except Exception as e:
            self.log_message(f"Critical error during processing: {str(e)}", "error")
            
        finally:
            # Re-enable button
            self.root.after(0, self.finish_processing)
    
    def save_results(self, df: pd.DataFrame, folder_path: str):
        """Save results to file"""
        try:
            output_format = self.output_format.get()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            if output_format == "excel":
                output_file = os.path.join(folder_path, f"candidate_extraction_results_{timestamp}.xlsx")
                df.to_excel(output_file, index=False)
            elif output_format == "csv":
                output_file = os.path.join(folder_path, f"candidate_extraction_results_{timestamp}.csv")
                df.to_csv(output_file, index=False)
            elif output_format == "json":
                output_file = os.path.join(folder_path, f"candidate_extraction_results_{timestamp}.json")
                df.to_json(output_file, orient='records', indent=2)
            
            self.log_message(f"Results saved to: {output_file}", "success")
            
        except Exception as e:
            self.log_message(f"Error saving results: {str(e)}", "error")
    
    def update_results_display(self, df: pd.DataFrame):
        """Update the results tree view"""
        # Configure columns
        columns = ['Filename', 'Candidate Name', 'Email ID', 'SAP Skills', 'Total Experience (Years)']
        self.results_tree['columns'] = columns
        self.results_tree['show'] = 'headings'
        
        # Configure column headings and widths
        for col in columns:
            self.results_tree.heading(col, text=col)
            if col == 'Filename':
                self.results_tree.column(col, width=150)
            elif col == 'Candidate Name':
                self.results_tree.column(col, width=120)
            elif col == 'Email ID':
                self.results_tree.column(col, width=150)
            elif col == 'SAP Skills':
                self.results_tree.column(col, width=200)
            else:
                self.results_tree.column(col, width=100)
        
        # Insert data
        for _, row in df.iterrows():
            values = [
                row['Filename'],
                row['Candidate Name'],
                row['Email ID'],
                row['SAP Skills'][:50] + "..." if len(str(row['SAP Skills'])) > 50 else row['SAP Skills'],
                row['Total Experience (Years)']
            ]
            self.results_tree.insert('', tk.END, values=values)
    
    def clear_results_tree(self):
        """Clear the results tree"""
        for item in self.results_tree.get_children():
            self.results_tree.delete(item)
    
    def show_processing_summary(self, df: pd.DataFrame):
        """Show processing summary"""
        total_files = len(df)
        successful_files = len(df[~df['Candidate Name'].str.contains('Error|Failed')])
        failed_files = total_files - successful_files
        success_rate = (successful_files / total_files) * 100 if total_files > 0 else 0
        
        self.log_message("=== Processing Summary ===", "info")
        self.log_message(f"Total files processed: {total_files}", "info")
        self.log_message(f"Successfully processed: {successful_files}", "success")
        self.log_message(f"Failed to process: {failed_files}", "error" if failed_files > 0 else "info")
        self.log_message(f"Success rate: {success_rate:.1f}%", "success")
        
        if failed_files > 0:
            self.log_message("Failed files:", "warning")
            failed_df = df[df['Candidate Name'].str.contains('Error|Failed')]
            for _, row in failed_df.iterrows():
                self.log_message(f"  - {row['Filename']}", "warning")
    
    def export_summary(self):
        """Export processing summary to file"""
        if not hasattr(self, 'extractor'):
            messagebox.showwarning("Warning", "No processing data available")
            return
        
        try:
            summary = self.extractor.get_processing_summary()
            
            # Save to JSON file
            folder = self.folder_path.get() or "."
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            summary_file = os.path.join(folder, f"processing_summary_{timestamp}.json")
            
            with open(summary_file, 'w') as f:
                json.dump(summary, f, indent=2)
            
            self.log_message(f"Processing summary exported to: {summary_file}", "success")
            messagebox.showinfo("Export Complete", f"Summary exported to:\n{summary_file}")
            
        except Exception as e:
            self.log_message(f"Error exporting summary: {str(e)}", "error")
            messagebox.showerror("Export Error", f"Failed to export summary: {str(e)}")
    
    def update_status(self, message: str):
        """Update status label"""
        self.root.after(0, lambda: self.status_label.config(text=message))
    
    def finish_processing(self):
        """Re-enable UI after processing"""
        self.process_btn.config(state="normal")
        self.processing = False
        self.update_status("Processing complete")

def main():
    root = tk.Tk()
    app = EnhancedCandidateExtractorGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()

Py Code:-

import os
import re
import pandas as pd
import logging
from typing import Dict, List, Tuple, Optional, Union
import json
from pathlib import Path
import spacy
from datetime import datetime
import traceback

# Import for file processing
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
except ImportError:
    Document = None

try:
    import textract
except ImportError:
    textract = None

class EnhancedCandidateExtractor:
    """Enhanced candidate extractor with multi-algorithm approach and detailed logging"""
    
    def __init__(self, log_level=logging.INFO):
        # Setup logging
        self.logger = logging.getLogger('CandidateExtractor')
        self.logger.setLevel(log_level)
        
        # Clear existing handlers
        for handler in self.logger.handlers[:]:
            self.logger.removeHandler(handler)
        
        # Create console handler
        handler = logging.StreamHandler()
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        handler.setFormatter(formatter)
        self.logger.addHandler(handler)
        
        # Initialize processing stats
        self.reset_stats()
        
        # Load NLP model
        self._load_nlp_model()
        
        # Define comprehensive patterns
        self._initialize_patterns()
    
    def reset_stats(self):
        """Reset processing statistics"""
        self.stats = {
            'total_files': 0,
            'processed_files': 0,
            'failed_files': [],
            'warnings': [],
            'processing_log': []
        }
    
    def _load_nlp_model(self):
        """Load spacy NLP model"""
        try:
            self.nlp = spacy.load("en_core_web_sm")
            self.logger.info("Successfully loaded spaCy model")
        except Exception as e:
            self.logger.warning(f"Could not load spaCy model: {e}")
            self.nlp = None
    
    def _initialize_patterns(self):
        """Initialize comprehensive extraction patterns"""
        # Email patterns
        self.email_patterns = [
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            r'\b[A-Za-z0-9._%+-]+\s*@\s*[A-Za-z0-9.-]+\s*\.\s*[A-Z|a-z]{2,}\b'
        ]
        
        # Experience patterns from filename
        self.filename_exp_patterns = [
            r'[\[\(](\d+)y[_\s]*(\d+)m[\]\)]',  # [6y_6m] or (6y_6m)
            r'[\[\(](\d+)yrs[_\s]*(\d+)months?[\]\)]',  # [6yrs_6months]
            r'[\[\(](\d+)years?[_\s]*(\d+)months?[\]\)]',  # [6years_6months]
            r'[\[\(](\d+)y(\d+)m[\]\)]',  # [6y6m]
            r'[\[\(](\d+)[_\s]*y[_\s]*(\d+)[_\s]*m[\]\)]',  # [6_y_6_m]
            r'(\d+)y[_\s]*(\d+)m',  # 6y_6m without brackets
            r'(\d+)[_\s]?yrs?[_\s]*(\d+)[_\s]?months?',  # 6yrs6months
        ]
        
        # Experience patterns from text
        self.text_exp_patterns = [
            r'(?:overall|total|having)\s*(\d+(?:\.\d+)?)\+?\s*years?',
            r'(\d+(?:\.\d+)?)\+?\s*years?\s*(?:of\s*)?(?:overall|total|work|professional|relevant)?\s*experience',
            r'experience\s*(?:of\s*)?(\d+(?:\.\d+)?)\+?\s*years?',
            r'(\d+(?:\.\d+)?)\+?\s*years?\s*experience',
            r'(\d+(?:\.\d+)?)\+?\s*years?\s*(?:into|in)',
            r'(\d+)\.\d+\s*years?',  # Decimal years like 5.8 years
        ]
        
        # SAP skills patterns
        self.sap_patterns = [
            r'SAP\s+(?:ERP\s+)?(?:FICO|FI|CO|S/4\s*HANA|ECC|MM|SD|PM|PP|HR|BASIS|ABAP|CRM|SRM|BW|BI|MDG|HANA)',
            r'SAP\s+(?:Certified|Professional|Associate|Expert)',
            r'(?:General\s+Ledger|GL|Accounts\s+Payable|AP|Accounts\s+Receivable|AR|Asset\s+Accounting|AA)',
            r'(?:Cost\s+Element|CEA|Cost\s+Center|CCA|Profit\s+Center|PCA|COPA)',
            r'(?:Financial\s+Accounting|Management\s+Accounting|Controlling)',
            r'S/4\s*HANA',
            r'Central\s+Finance|CFIN',
            r'SAP\s+modules?',
            r'LSMW|BAPI|RFC|IDOC|SLT',
        ]
        
        # Name patterns for better extraction
        self.name_indicators = [
            r'^(?:Mr\.?|Ms\.?|Mrs\.?|Dr\.?)\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
            r'^([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
            r'Name\s*:?\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
        ]
        
        # Certification patterns
        self.cert_patterns = [
            r'SAP\s+Certified[^.]*',
            r'Certification[s]?\s*:?\s*([^\n]+)',
            r'(?:Certified|Certificate)\s+(?:in\s+)?([^\n,]+)',
            r'(?:DELF|IELTS|PMP|CISSP|CISA|FRM|CFA|ACCA)[^.\n]*',
        ]
    
    def log_process_step(self, step: str, details: str = "", level: str = "info"):
        """Log processing step with timestamp"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_entry = f"[{timestamp}] {step}"
        if details:
            log_entry += f" - {details}"
        
        self.stats['processing_log'].append(log_entry)
        
        if level == "error":
            self.logger.error(log_entry)
        elif level == "warning":
            self.logger.warning(log_entry)
        else:
            self.logger.info(log_entry)
    
    def extract_text_from_file(self, file_path: str) -> Tuple[str, bool]:
        """Extract text from various file formats with enhanced error handling"""
        file_path = Path(file_path)
        success = False
        text = ""
        
        self.log_process_step(f"Processing file", f"{file_path.name}")
        
        try:
            # Handle DOCX files
            if file_path.suffix.lower() == '.docx':
                if Document:
                    try:
                        doc = Document(file_path)
                        paragraphs = [para.text for para in doc.paragraphs]
                        
                        # Extract table content as well
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    paragraphs.append(cell.text)
                        
                        text = '\n'.join(paragraphs)
                        success = True
                        self.log_process_step("DOCX extraction", "Success")
                    except Exception as e:
                        self.log_process_step("DOCX extraction failed", str(e), "warning")
                        # Fallback to textract
                        if textract:
                            text = textract.process(file_path).decode('utf-8', errors='ignore')
                            success = True
                            self.log_process_step("Textract fallback", "Success")
                else:
                    self.log_process_step("DOCX library not available", "", "warning")
            
            # Handle DOC files
            elif file_path.suffix.lower() == '.doc':
                # Skip DOC files if textract is not available
                self.log_process_step("DOC file skipped", "Textract not available", "warning")
                return "DOC files not supported without textract library", False
            
            # Handle PDF files
            elif file_path.suffix.lower() == '.pdf':
                if pdfplumber:
                    try:
                        with pdfplumber.open(file_path) as pdf:
                            text = '\n'.join([page.extract_text() or '' for page in pdf.pages])
                        success = True
                        self.log_process_step("PDF extraction", "Success")
                    except Exception as e:
                        self.log_process_step("PDF extraction failed", str(e), "error")
                else:
                    self.log_process_step("PDFplumber not available", "", "error")
            
            else:
                self.log_process_step("Unsupported file format", file_path.suffix, "error")
        
        except Exception as e:
            self.log_process_step("File processing error", str(e), "error")
            text = ""
            success = False
        
        return text.strip(), success
    
    def extract_name_multi_algo(self, text: str, filename: str) -> str:
        """Extract candidate name using multiple algorithms"""
        self.log_process_step("Extracting name", "Using multi-algorithm approach")
        
        # Algorithm 1: Extract from filename
        filename_stem = Path(filename).stem
        # Remove common prefixes and clean filename
        filename_clean = re.sub(r'^(?:Naukri_|Resume_|CV_)', '', filename_stem, flags=re.IGNORECASE)
        filename_clean = re.sub(r'\[.*?\]|\(.*?\)', '', filename_clean)
        filename_clean = re.sub(r'[_-]', ' ', filename_clean).strip()
        
        if len(filename_clean.split()) >= 2 and all(part.isalpha() for part in filename_clean.split()):
            self.log_process_step("Name from filename", filename_clean)
            return filename_clean.title()
        
        # Algorithm 2: Extract from first lines
        lines = text.split('\n')[:10]
        for line in lines:
            line = line.strip()
            if not line or len(line) < 3:
                continue
                
            # Remove special characters and formatting
            clean_line = re.sub(r'[#*_|\-]+', '', line).strip()
            
            # Check for name patterns
            for pattern in self.name_indicators:
                match = re.search(pattern, clean_line)
                if match:
                    name = match.group(1).strip()
                    if len(name.split()) >= 2:
                        self.log_process_step("Name from pattern", name)
                        return name
            
            # Simple name detection: 2-4 words, all starting with capital
            words = clean_line.split()
            if 2 <= len(words) <= 4:
                if all(word[0].isupper() and word[1:].islower() for word in words if word.isalpha()):
                    if not any(keyword in clean_line.lower() for keyword in ['email', 'phone', 'contact', 'summary', 'profile', 'experience']):
                        self.log_process_step("Name from simple detection", clean_line)
                        return clean_line
        
        # Algorithm 3: NLP-based extraction
        if self.nlp:
            doc = self.nlp(text[:1000])  # First 1000 characters
            persons = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]
            if persons:
                # Filter out common false positives
                for person in persons:
                    if len(person.split()) >= 2 and not any(keyword in person.lower() for keyword in ['email', 'phone', 'contact']):
                        self.log_process_step("Name from NLP", person)
                        return person
        
        # Algorithm 4: Fallback to filename without extension
        if filename_clean:
            self.log_process_step("Name fallback", filename_clean)
            return filename_clean.title()
        
        self.log_process_step("Name extraction", "Failed to extract", "warning")
        return "Name not found"
    
    def extract_email_multi_algo(self, text: str) -> str:
        """Extract email using multiple patterns"""
        self.log_process_step("Extracting email", "Using multiple patterns")
        
        for pattern in self.email_patterns:
            matches = re.findall(pattern, text)
            if matches:
                # Filter out generic emails
                for email in matches:
                    if not any(generic in email.lower() for generic in ['example', 'test', 'sample', 'domain']):
                        self.log_process_step("Email found", email)
                        return email
        
        self.log_process_step("Email extraction", "Not found", "warning")
        return "Email not found"
    
    def extract_sap_skills_multi_algo(self, text: str) -> str:
        """Extract SAP skills using comprehensive pattern matching"""
        self.log_process_step("Extracting SAP skills", "Using comprehensive patterns")
        
        found_skills = set()
        text_lower = text.lower()
        
        # Algorithm 1: Direct pattern matching
        for pattern in self.sap_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                found_skills.add(match.strip())
        
        # Algorithm 2: Context-based extraction
        sap_keywords = [
            'SAP FICO', 'SAP FI', 'SAP CO', 'SAP MM', 'SAP SD', 'SAP PM', 'SAP PP', 
            'SAP HR', 'SAP BASIS', 'SAP ABAP', 'S/4 HANA', 'SAP ECC', 'SAP CRM',
            'SAP SRM', 'SAP BW', 'SAP BI', 'SAP MDG', 'Central Finance', 'CFIN',
            'General Ledger', 'Accounts Payable', 'Accounts Receivable',
            'Asset Accounting', 'Cost Center', 'Profit Center', 'COPA'
        ]
        
        for keyword in sap_keywords:
            if keyword.lower() in text_lower:
                found_skills.add(keyword)
        
        # Algorithm 3: Extract from sections
        sections = re.split(r'\n(?=[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s*:)', text)
        for section in sections:
            if any(header in section.lower() for header in ['skill', 'experience', 'expertise', 'competenc']):
                for pattern in self.sap_patterns:
                    matches = re.findall(pattern, section, re.IGNORECASE)
                    for match in matches:
                        found_skills.add(match.strip())
        
        if found_skills:
            skills_text = ', '.join(sorted(found_skills))
            self.log_process_step("SAP skills found", f"{len(found_skills)} skills")
            return skills_text
        
        self.log_process_step("SAP skills", "Not found", "warning")
        return "No SAP skills found"
    
    def extract_experience_multi_algo(self, text: str, filename: str) -> float:
        """Extract total experience using multiple algorithms with priority"""
        self.log_process_step("Extracting experience", "Using priority-based algorithms")
        
        # Algorithm 1: Filename parsing (highest priority)
        experience = self._extract_experience_from_filename(filename)
        if experience is not None:
            self.log_process_step("Experience from filename", f"{experience} years")
            return experience
        
        # Algorithm 2: Text pattern matching
        experience = self._extract_experience_from_text(text)
        if experience is not None:
            self.log_process_step("Experience from text", f"{experience} years")
            return experience
        
        self.log_process_step("Experience extraction", "Failed", "warning")
        return 0.0
    
    def _extract_experience_from_filename(self, filename: str) -> Optional[float]:
        """Extract experience from filename using comprehensive patterns"""
        for pattern in self.filename_exp_patterns:
            match = re.search(pattern, filename, re.IGNORECASE)
            if match:
                try:
                    years = int(match.group(1))
                    months = int(match.group(2))
                    total_years = years + months / 12.0
                    return round(total_years, 1)
                except (ValueError, IndexError):
                    continue
        return None
    
    def _extract_experience_from_text(self, text: str) -> Optional[float]:
        """Extract experience from text using multiple patterns"""
        for pattern in self.text_exp_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                try:
                    # Take the highest reasonable experience value
                    experiences = [float(match) for match in matches if float(match) <= 50]
                    if experiences:
                        return max(experiences)
                except ValueError:
                    continue
        return None
    
    def extract_summary_multi_algo(self, text: str) -> str:
        """Extract professional summary excluding personal information"""
        self.log_process_step("Extracting summary", "Multi-algorithm approach")
        
        # Look for summary sections
        summary_patterns = [
            r'(?:PROFESSIONAL\s+)?SUMMARY\s*:?\s*(.*?)(?=\n(?:[A-Z][A-Z\s]*:|$))',
            r'PROFILE\s+SUMMARY\s*:?\s*(.*?)(?=\n(?:[A-Z][A-Z\s]*:|$))',
            r'OBJECTIVE\s*:?\s*(.*?)(?=\n(?:[A-Z][A-Z\s]*:|$))',
            r'ABOUT\s*:?\s*(.*?)(?=\n(?:[A-Z][A-Z\s]*:|$))',
        ]
        
        for pattern in summary_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                summary = match.group(1).strip()
                # Clean up the summary
                summary = re.sub(r'\s+', ' ', summary)  # Normalize whitespace
                summary = summary[:500]  # Limit length
                
                # Remove personal information
                summary = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '[EMAIL]', summary)
                summary = re.sub(r'\b(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})\b', '[PHONE]', summary)
                
                if len(summary) > 50:
                    self.log_process_step("Summary extracted", f"{len(summary)} characters")
                    return summary
        
        # Fallback: Extract first meaningful paragraph
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if 100 < len(para) < 800 and not any(keyword in para.lower() for keyword in ['contact', 'email', 'phone', 'address']):
                para = re.sub(r'\s+', ' ', para)
                self.log_process_step("Summary from paragraph", f"{len(para)} characters")
                return para[:500]
        
        self.log_process_step("Summary", "Basic summary generated")
        return "Professional with relevant experience in the field"
    
    def extract_certifications_multi_algo(self, text: str) -> str:
        """Extract certifications using comprehensive patterns"""
        self.log_process_step("Extracting certifications", "Comprehensive pattern matching")
        
        found_certs = set()
        
        # Look for certification sections
        cert_section_pattern = r'CERTIFICATION[S]?\s*:?\s*(.*?)(?=\n(?:[A-Z][A-Z\s]*:|$))'
        match = re.search(cert_section_pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            cert_text = match.group(1)
            # Clean up certification text
            cert_lines = cert_text.split('\n')
            for line in cert_lines:
                line = line.strip()
                if line and len(line) > 10:
                    # Remove dates and locations
                    cleaned = re.sub(r'\b\d{4}\b', '', line)  # Remove years
                    cleaned = re.sub(r'\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b', '', cleaned)  # Remove dates
                    cleaned = re.sub(r'\([^)]*\)', '', cleaned)  # Remove parentheses content
                    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
                    if cleaned and len(cleaned) > 5:
                        found_certs.add(cleaned)
        
        # Pattern-based extraction
        for pattern in self.cert_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    match = match[0] if match[0] else match[1] if len(match) > 1 else ''
                
                # Clean up the certification
                cleaned = re.sub(r'\b\d{4}\b', '', match)
                cleaned = re.sub(r'\([^)]*\)', '', cleaned)
                cleaned = re.sub(r'\s+', ' ', cleaned).strip()
                
                if cleaned and len(cleaned) > 5:
                    found_certs.add(cleaned)
        
        if found_certs:
            certs_text = ', '.join(sorted(found_certs))
            self.log_process_step("Certifications found", f"{len(found_certs)} certifications")
            return certs_text
        
        self.log_process_step("Certifications", "Not found", "warning")
        return "No certifications found"
    
    def process_single_file(self, file_path: str) -> Dict:
        """Process a single file and extract all information"""
        self.stats['total_files'] += 1
        file_path = Path(file_path)
        
        self.log_process_step("Starting file processing", file_path.name)
        
        try:
            # Extract text
            text, success = self.extract_text_from_file(file_path)
            
            if not success or not text.strip():
                error_msg = f"Failed to extract text from {file_path.name}"
                self.stats['failed_files'].append(f"{file_path.name}: {error_msg}")
                self.log_process_step("File processing failed", error_msg, "error")
                return {
                    'Filename': file_path.name,
                    'Candidate Name': 'Extraction Failed',
                    'Email ID': 'Extraction Failed',
                    'SAP Skills': 'Extraction Failed',
                    'Total Experience (Years)': 0,
                    'Summary': 'Failed to process file',
                    'Certifications': 'Extraction Failed'
                }
            
            # Extract all fields using multi-algorithm approach
            candidate_name = self.extract_name_multi_algo(text, file_path.name)
            email = self.extract_email_multi_algo(text)
            sap_skills = self.extract_sap_skills_multi_algo(text)
            experience = self.extract_experience_multi_algo(text, file_path.name)
            summary = self.extract_summary_multi_algo(text)
            certifications = self.extract_certifications_multi_algo(text)
            
            self.stats['processed_files'] += 1
            self.log_process_step("File processing completed", file_path.name)
            
            return {
                'Filename': file_path.name,
                'Candidate Name': candidate_name,
                'Email ID': email,
                'SAP Skills': sap_skills,
                'Total Experience (Years)': experience,
                'Summary': summary,
                'Certifications': certifications
            }
        
        except Exception as e:
            error_msg = f"Unexpected error processing {file_path.name}: {str(e)}"
            self.stats['failed_files'].append(f"{file_path.name}: {error_msg}")
            self.log_process_step("Unexpected error", error_msg, "error")
            
            return {
                'Filename': file_path.name,
                'Candidate Name': 'Processing Error',
                'Email ID': 'Processing Error',
                'SAP Skills': 'Processing Error',
                'Total Experience (Years)': 0,
                'Summary': f'Error: {str(e)}',
                'Certifications': 'Processing Error'
            }
    
    def process_files(self, file_paths: List[str]) -> pd.DataFrame:
        """Process multiple files and return results as DataFrame"""
        self.reset_stats()
        self.log_process_step("Starting batch processing", f"{len(file_paths)} files")
        
        results = []
        
        for file_path in file_paths:
            result = self.process_single_file(file_path)
            results.append(result)
        
        # Create summary statistics
        success_rate = (self.stats['processed_files'] / self.stats['total_files']) * 100 if self.stats['total_files'] > 0 else 0
        
        self.log_process_step("Batch processing completed", 
                            f"Success: {self.stats['processed_files']}/{self.stats['total_files']} ({success_rate:.1f}%)")
        
        if self.stats['failed_files']:
            self.log_process_step("Failed files", f"{len(self.stats['failed_files'])} files failed", "warning")
            for failed_file in self.stats['failed_files']:
                self.log_process_step("Failed file details", failed_file, "warning")
        
        return pd.DataFrame(results)
    
    def get_processing_summary(self) -> Dict:
        """Get detailed processing summary"""
        return {
            'total_files': self.stats['total_files'],
            'processed_successfully': self.stats['processed_files'],
            'failed_files': self.stats['failed_files'],
            'success_rate': (self.stats['processed_files'] / self.stats['total_files']) * 100 if self.stats['total_files'] > 0 else 0,
            'processing_log': self.stats['processing_log']
        }
    
    def export_results(self, df: pd.DataFrame, output_path: str, format_type: str = 'excel'):
        """Export results to specified format"""
        self.log_process_step("Exporting results", f"Format: {format_type}, Path: {output_path}")
        
        try:
            if format_type.lower() == 'excel':
                df.to_excel(output_path, index=False)
            elif format_type.lower() == 'csv':
                df.to_csv(output_path, index=False)
            elif format_type.lower() == 'json':
                df.to_json(output_path, orient='records', indent=2)
            
            self.log_process_step("Export completed", output_path)
        except Exception as e:
            self.log_process_step("Export failed", str(e), "error")

if __name__ == "__main__":
    # Example usage
    extractor = EnhancedCandidateExtractor()
    
    # Test with sample files
    sample_files = [
        "user_input_files/GouriRSalunkhe.docx",
        "user_input_files/Sintukumarsingh.docx",
        "user_input_files/Pawan Deshwal.docx",
        "user_input_files/Naukri_TNataraja[5y_8m].docx",
        "user_input_files/Naukri_Kishorekumar[8y_0m] (1).docx"
    ]
    
    # Process files
    results_df = extractor.process_files(sample_files)
    
    # Export results
    extractor.export_results(results_df, "enhanced_results.xlsx", "excel")
    
    # Print processing summary
    summary = extractor.get_processing_summary()
    print(f"\nProcessing Summary:")
    print(f"Total files: {summary['total_files']}")
    print(f"Successfully processed: {summary['processed_successfully']}")
    print(f"Failed files: {len(summary['failed_files'])}")
    print(f"Success rate: {summary['success_rate']:.1f}%")
    
    if summary['failed_files']:
        print(f"\nFailed files:")
        for failed in summary['failed_files']:
            print(f"  - {failed}")
    
    print(f"\nProcessing Log:")
    for log_entry in summary['processing_log'][-10:]:  # Show last 10 entries
        print(f"  {log_entry}")
